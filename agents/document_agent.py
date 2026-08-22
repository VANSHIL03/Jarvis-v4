"""
JARVIS v4 - Document Understanding Agent

Section 18: scan a document, pull its text out, summarise it, turn a scan into
something editable, and write a PDF back. PDF, PNG, JPG, JPEG, DOCX and TXT are
all handled through one extraction path so every action behaves the same way
regardless of which format the user happened to hand over.

Nothing is reimplemented here: PDFs go through automation/office.py, images
through the OCR in ai/vision.py, summaries through the local LLM, and PDF
writing through OfficeAutomation.create_pdf.
"""

from pathlib import Path
from typing import Any, Dict, List, Optional

from agents.base_agent import BaseAgent
from ai.llm_client import LocalLLMClient
from automation.file_manager import FileManager
from automation.office import OfficeAutomation
from utils.logger import logger

#: Extensions this agent knows how to read.
SUPPORTED_EXTENSIONS = {".pdf", ".png", ".jpg", ".jpeg", ".docx", ".txt", ".md"}


class DocumentAgent(BaseAgent):
    def __init__(
        self,
        llm_client: Optional[LocalLLMClient] = None,
        office_auto: Optional[OfficeAutomation] = None,
        vision_analyzer: Any = None,
        file_manager: Optional[FileManager] = None,
    ):
        self.llm = llm_client or LocalLLMClient()
        self.office = office_auto or OfficeAutomation()
        self.files = file_manager or FileManager()
        self._vision = vision_analyzer

    @property
    def agent_name(self) -> str:
        return "document_agent"

    @property
    def description(self) -> str:
        return (
            "Reads and summarises PDF, DOCX, TXT and scanned image documents, "
            "converts scans to editable text, and writes PDFs."
        )

    # ------------------------------------------------------------- helpers
    def _vision_analyzer(self):
        """Lazily resolves an OCR-capable analyzer for image documents."""
        if self._vision is None or not hasattr(self._vision, "extract_text"):
            from ai.vision import VisionAnalyzer
            self._vision = VisionAnalyzer(llm_client=self.llm)
        return self._vision

    def _resolve(self, raw_path: str) -> Path:
        return self.files.resolve_path(raw_path)

    def _read_docx(self, path: Path) -> str:
        """Extracts paragraphs and table cells from a .docx."""
        try:
            import docx
        except ImportError:
            logger.warning("python-docx not installed; cannot read .docx files.")
            return ""
        try:
            document = docx.Document(str(path))
            chunks: List[str] = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        chunks.append(" | ".join(cells))
            return "\n".join(chunks).strip()
        except Exception as e:
            logger.error(f"Failed to read '{path}': {e}")
            return ""

    def _extract(self, raw_path: str) -> Dict[str, Any]:
        """
        Pulls plain text out of any supported document.

        Returns a status dict rather than raising, so every caller in this agent
        can hand the same failure straight back to the user.
        """
        if not str(raw_path or "").strip():
            return {
                "status": "error",
                "message": "No document path given.",
                "speech_reply": "Sir, kaunsa document padhna hai wo bataiye.",
            }

        path = self._resolve(raw_path)
        if not path.exists():
            return {
                "status": "not_found",
                "path": str(path),
                "message": f"'{path}' not found.",
                "speech_reply": f"Sir, '{path.name}' mila hi nahi.",
            }

        suffix = path.suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            return {
                "status": "error",
                "path": str(path),
                "message": (
                    f"'{suffix or 'no extension'}' is not supported. "
                    f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}."
                ),
                "speech_reply": (
                    f"Sir, main {suffix or 'is'} format nahi padh sakta. "
                    "PDF, DOCX, TXT, PNG ya JPG dijiye."
                ),
            }

        if suffix == ".pdf":
            text, kind = self.office.read_pdf(str(path)), "pdf"
        elif suffix == ".docx":
            text, kind = self._read_docx(path), "docx"
        elif suffix in (".txt", ".md"):
            try:
                text = path.read_text(encoding="utf-8", errors="replace")
            except Exception as e:
                return {"status": "error", "path": str(path), "message": str(e)}
            kind = "text"
        else:
            text, kind = self._vision_analyzer().extract_text(str(path)), "image"

        text = (text or "").strip()
        if not text:
            return {
                "status": "error",
                "path": str(path),
                "kind": kind,
                "message": (
                    "No text could be extracted. A scanned PDF needs OCR, and an "
                    "image needs Tesseract installed."
                ),
                "speech_reply": (
                    f"Sir, '{path.name}' me se koi text nahi nikla - lagta hai ye "
                    "scan hai jiske liye OCR available nahi hai."
                ),
            }

        return {
            "status": "success",
            "path": str(path),
            "kind": kind,
            "text": text,
            "chars": len(text),
        }

    # ------------------------------------------------------------ dispatch
    async def execute_task(self, action: str, params: Dict[str, Any]) -> Dict[str, Any]:
        action = (action or "").lower().strip()
        raw_path = (
            params.get("path")
            or params.get("file_path")
            or params.get("document")
            or params.get("image_path")
            or ""
        )

        if action in ("scan_document", "scan", "read_document"):
            found = self._extract(raw_path)
            if found["status"] != "success":
                return found
            text = found["text"]
            preview = " ".join(text.split())[:400]
            words = len(text.split())
            found.update({
                "preview": preview,
                "words": words,
                "speech_reply": (
                    f"Ji Sir, '{Path(found['path']).name}' scan ho gaya - "
                    f"{words} words mile. Shuruaat: {preview[:200]}"
                ),
            })
            return found

        if action in ("extract_text", "get_text", "ocr_document"):
            found = self._extract(raw_path)
            if found["status"] != "success":
                return found
            found["speech_reply"] = (
                f"Ji Sir, '{Path(found['path']).name}' se {found['chars']} characters "
                "ka text nikal liya hai."
            )
            return found

        if action in ("summarize_document", "summarise_document", "summarize", "summarize_pdf"):
            found = self._extract(raw_path)
            if found["status"] != "success":
                return found

            instructions = params.get("prompt") or params.get("user_prompt") or ""
            body = found["text"][:12000]
            system_prompt = (
                "You are JARVIS, summarising a document for your user. Be accurate "
                "and concise. Use short bullet points. Never invent details that are "
                "not in the text."
            )
            prompt = (
                f"Document: {Path(found['path']).name}\n\n"
                f"Content:\n```\n{body}\n```\n\n"
                + (f"Extra instructions: {instructions}\n\n" if instructions else "")
                + "Summarise the key points."
            )
            try:
                summary = await self.llm.generate_response(prompt=prompt, system_prompt=system_prompt)
            except Exception as e:
                logger.error(f"Summarisation failed: {e}")
                return {
                    "status": "error",
                    "path": found["path"],
                    "message": f"LLM unavailable: {e}",
                    "speech_reply": "Sir, summary banane ke liye local model available nahi hai.",
                }

            summary = (summary or "").strip()
            if not summary:
                return {
                    "status": "error",
                    "path": found["path"],
                    "message": "Model returned an empty summary.",
                    "speech_reply": "Sir, model se koi summary nahi aayi.",
                }

            save_to = params.get("save_to") or ""
            saved: Dict[str, Any] = {}
            if save_to:
                saved = self.office.create_pdf(
                    str(self._resolve(save_to)),
                    title=f"Summary - {Path(found['path']).name}",
                    body=summary,
                )

            return {
                "status": "success",
                "path": found["path"],
                "kind": found["kind"],
                "summary": summary,
                "message": summary,
                "saved": saved,
                "speech_reply": f"Sir, is document ka summary ye hai: {summary[:400]}",
            }

        if action in ("to_editable_text", "make_editable", "convert_to_word", "to_word"):
            found = self._extract(raw_path)
            if found["status"] != "success":
                return found

            source = Path(found["path"])
            out_raw = params.get("save_to") or params.get("output_path") or ""
            out_path = (
                self._resolve(out_raw) if out_raw
                else source.with_name(f"{source.stem}_editable.docx")
            )
            paragraphs = [p.strip() for p in found["text"].split("\n") if p.strip()]
            ok = self.office.create_word_document(str(out_path), source.stem, paragraphs)
            if not ok:
                return {
                    "status": "error",
                    "path": found["path"],
                    "message": f"Could not write '{out_path}'.",
                    "speech_reply": "Sir, editable file banane me problem aa gayi.",
                }
            return {
                "status": "success",
                "source": found["path"],
                "path": str(out_path),
                "paragraphs": len(paragraphs),
                "speech_reply": f"Ji Sir, editable version bana diya: {out_path}",
            }

        if action in ("create_pdf", "make_pdf", "write_pdf", "save_as_pdf"):
            target_raw = (
                params.get("save_to")
                or params.get("output_path")
                or params.get("path")
                or ""
            )
            title = params.get("title", "")
            body = params.get("body") or params.get("content") or params.get("text") or ""

            source = params.get("source") or params.get("from_document") or ""
            if source and not body:
                found = self._extract(source)
                if found["status"] != "success":
                    return found
                body = found["text"]
                title = title or Path(found["path"]).stem

            if not target_raw:
                stem = "".join(c for c in (title or "JARVIS_Document") if c.isalnum() or c in " _-").strip()
                target_raw = f"{stem or 'JARVIS_Document'}.pdf"
            target = self._resolve(target_raw)
            if target.suffix.lower() != ".pdf":
                target = target.with_suffix(".pdf")

            if not str(body).strip():
                return {
                    "status": "error",
                    "message": "No content supplied for the PDF.",
                    "speech_reply": "Sir, PDF me kya likhna hai wo bataiye.",
                }

            res = self.office.create_pdf(str(target), title=title, body=str(body))
            if res.get("status") == "success":
                res.setdefault(
                    "speech_reply", f"Ji Sir, PDF ready hai: {res.get('path', target)}"
                )
            else:
                res.setdefault("speech_reply", "Sir, PDF banane me problem aa gayi.")
            return res

        return {"status": "error", "message": f"Unknown document action: '{action}'"}
